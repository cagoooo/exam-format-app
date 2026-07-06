from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import time
import uuid
from copy import deepcopy
from dataclasses import dataclass, replace
from io import BytesIO
from pathlib import Path

from flask import Flask, jsonify, render_template, request, send_file
from werkzeug.utils import secure_filename

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips


ROOT = Path(__file__).resolve().parent
DATA_DIR = Path(os.environ.get("EXAM_FORMAT_DATA_DIR", ROOT))
UPLOAD_DIR = DATA_DIR / "uploads"
GENERATED_DIR = DATA_DIR / "generated"
CONFIG_DIR = ROOT / "config"
PROFILE_PATH = CONFIG_DIR / "profiles.json"
VERSION_PATH = ROOT / "version.json"
ALLOWED_EXT = {".doc", ".docx"}
MAX_UPLOAD_BYTES = 25 * 1024 * 1024
JOB_TTL_SECONDS = int(os.environ.get("JOB_TTL_SECONDS", "1800"))
SCHOOL_NAME = "桃園市龍潭區石門國民小學"
APP_TITLE = "考卷格式自動校正系統"
APP_DESCRIPTION = "固定版面、頁數檢查、Word 與 PDF 輸出的考卷格式標準化工具。"
ASSET_VERSION = "20260706-og1"


@dataclass(frozen=True)
class ExamProfile:
    key: str
    label: str
    width: int
    height: int
    margins: dict[str, int]
    columns: int = 2
    column_space: int = 425
    column_separator: bool = True
    header_footer: tuple[int, int] = (851, 992)
    title_size: int = 16
    body_size: int = 11
    question_size: int = 11
    line_spacing: int = 15
    table_font_size: int = 9
    font: str = "標楷體"


def load_profiles() -> dict[str, ExamProfile]:
    raw = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
    profiles: dict[str, ExamProfile] = {}
    for key, value in raw.items():
        profiles[key] = ExamProfile(
            key=key,
            label=value["label"],
            width=value["width"],
            height=value["height"],
            margins=value["margins"],
            columns=value.get("columns", 2),
            column_space=value.get("column_space", 425),
            column_separator=value.get("column_separator", True),
            header_footer=tuple(value.get("header_footer", [851, 992])),
            title_size=value.get("title_size", 16),
            body_size=value.get("body_size", 11),
            question_size=value.get("question_size", 11),
            line_spacing=value.get("line_spacing", 15),
            table_font_size=value.get("table_font_size", 9),
            font=value.get("font", "標楷體"),
        )
    return profiles


def load_app_version() -> str:
    try:
        return json.loads(VERSION_PATH.read_text(encoding="utf-8")).get("version", ASSET_VERSION)
    except Exception:
        return ASSET_VERSION


PROFILES: dict[str, ExamProfile] = load_profiles()


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_BYTES


@app.after_request
def add_cors_headers(response):
    origin = request.headers.get("Origin")
    allowed_origins = {
        "https://cagoooo.github.io",
        "http://127.0.0.1:5127",
        "http://localhost:5127",
    }
    if origin in allowed_origins or os.environ.get("ALLOW_ANY_ORIGIN") == "1":
        response.headers["Access-Control-Allow-Origin"] = origin or "*"
        response.headers["Vary"] = "Origin"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type"
    return response


def ensure_work_dirs() -> None:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)


def cleanup_expired_jobs() -> None:
    now = time.time()
    for root in (UPLOAD_DIR, GENERATED_DIR):
        if not root.exists():
            continue
        for job_dir in root.iterdir():
            if not job_dir.is_dir():
                continue
            try:
                age = now - job_dir.stat().st_mtime
            except FileNotFoundError:
                continue
            if age > JOB_TTL_SECONDS:
                shutil.rmtree(job_dir, ignore_errors=True)


def create_job_dirs() -> tuple[str, Path, Path]:
    ensure_work_dirs()
    cleanup_expired_jobs()
    job_id = uuid.uuid4().hex
    upload_dir = UPLOAD_DIR / job_id
    generated_dir = GENERATED_DIR / job_id
    upload_dir.mkdir(parents=True, exist_ok=False)
    generated_dir.mkdir(parents=True, exist_ok=False)
    return job_id, upload_dir, generated_dir


def resolve_job_file(root: Path, job_id: str, filename: str) -> Path | None:
    if not re.fullmatch(r"[0-9a-f]{32}", job_id):
        return None
    job_dir = (root / job_id).resolve()
    root_resolved = root.resolve()
    if job_dir.parent != root_resolved or not job_dir.exists():
        return None
    if time.time() - job_dir.stat().st_mtime > JOB_TTL_SECONDS:
        shutil.rmtree(job_dir, ignore_errors=True)
        return None
    path = (job_dir / filename).resolve()
    if path.parent != job_dir or not path.exists():
        return None
    os.utime(job_dir, None)
    return path


def public_file_path(kind: str, job_id: str, filename: str) -> str:
    return f"/{kind}/{job_id}/{filename}"


def error_response(message: str, status: int = 400, code: str = "REQUEST_ERROR", details: list[str] | None = None):
    return jsonify({"ok": False, "error": message, "code": code, "details": details or []}), status


def allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXT


def safe_upload_name(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    stem = secure_filename(Path(filename).stem) or "exam"
    return f"{stem}-{uuid.uuid4().hex[:8]}{ext}"


def find_soffice() -> str | None:
    for command in ("soffice", "libreoffice"):
        resolved = shutil.which(command)
        if resolved:
            return resolved
    return None


def convert_with_libreoffice(path: Path, target_ext: str) -> Path:
    soffice = find_soffice()
    if not soffice:
        raise RuntimeError("線上轉檔需要 LibreOffice；目前執行環境找不到 soffice/libreoffice。")

    out_dir = path.parent
    before = {p.name for p in out_dir.glob(f"*{target_ext}")}
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                target_ext.lstrip("."),
                "--outdir",
                str(out_dir),
                str(path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120,
        )
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError("轉檔逾時。檔案可能太大、內容過於複雜，請先移除大型圖片或分成較小份考卷後再試。") from exc
    except subprocess.CalledProcessError as exc:
        stderr = exc.stderr.decode("utf-8", errors="ignore") if exc.stderr else ""
        raise RuntimeError(f"LibreOffice 轉檔失敗。請確認檔案沒有加密、損毀或仍被 Word 開啟。{stderr[:200]}") from exc
    expected = path.with_suffix(target_ext)
    if expected.exists():
        return expected
    after = [p for p in out_dir.glob(f"*{target_ext}") if p.name not in before]
    if after:
        return after[0]
    raise RuntimeError(f"LibreOffice 已執行，但沒有產生 {target_ext} 檔。")


def convert_doc_to_docx(path: Path) -> Path:
    out = path.with_suffix(".docx")
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(path))
        doc.SaveAs(str(out), FileFormat=16)
        doc.Close(False)
        word.Quit()
        return out
    except Exception:
        return convert_with_libreoffice(path, ".docx")


def iter_docx_blocks(path: Path) -> list[dict]:
    doc = Document(str(path))
    blocks: list[dict] = []

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, doc)
            text = normalize_text(paragraph.text)
            has_special_inline = bool(child.xpath(".//m:oMath | .//m:oMathPara | .//w:ruby"))
            runs = []
            images = []
            for run in paragraph.runs:
                run_text = run.text
                if run_text:
                    runs.append(
                        {
                            "text": run_text,
                            "bold": bool(run.bold),
                            "italic": bool(run.italic),
                            "underline": bool(run.underline),
                        }
                    )
                for blip in run._element.xpath(".//a:blip"):
                    rel_id = blip.get(qn("r:embed"))
                    if not rel_id:
                        continue
                    part = paragraph.part.related_parts.get(rel_id)
                    if not part:
                        continue
                    images.append(
                        {
                            "rel_id": rel_id,
                            "blob": part.blob,
                            "content_type": part.content_type,
                            "filename": Path(part.partname).name,
                        }
                    )
            if text or images or has_special_inline:
                blocks.append(
                    {
                        "type": "paragraph",
                        "text": text,
                        "runs": runs,
                        "images": images,
                        "xml": deepcopy(child) if has_special_inline else None,
                        "preserve_xml": has_special_inline,
                        "has_formula": bool(child.xpath(".//m:oMath | .//m:oMathPara")),
                        "has_ruby": bool(child.xpath(".//w:ruby")),
                    }
                )
        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [normalize_text(cell.text) for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                table_images = {}
                for blip in child.xpath(".//a:blip"):
                    rel_id = blip.get(qn("r:embed"))
                    if not rel_id:
                        continue
                    part = doc.part.related_parts.get(rel_id)
                    if part:
                        table_images[rel_id] = {"blob": part.blob, "content_type": part.content_type}
                blocks.append({"type": "table", "rows": rows, "xml": deepcopy(child), "images": table_images})

    return blocks


def preview_blocks(blocks: list[dict], limit: int = 8) -> list[str]:
    lines: list[str] = []
    for block in blocks:
        if block["type"] == "paragraph":
            lines.append(block["text"])
        elif block["type"] == "table":
            for row in block["rows"][:4]:
                text = " | ".join(cell for cell in row if cell)
                if text:
                    lines.append(text)
        if len(lines) >= limit:
            break
    return lines[:limit]


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def looks_like_question(text: str) -> bool:
    return bool(re.match(r"^([一二三四五六七八九十]+、|\d+[\.、)]|[（(]\d+[）)])", text.strip()))


def set_cell_text(cell, text: str, font: str, size: int, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run, font, size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, font: str, size: int) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_styled_runs(paragraph, runs: list[dict], fallback_text: str, font: str, size: int) -> None:
    if not runs:
        run = paragraph.add_run(fallback_text)
        set_run_font(run, font, size)
        return
    for source_run in runs:
        run = paragraph.add_run(source_run["text"])
        set_run_font(run, font, size)
        run.bold = source_run.get("bold", False)
        run.italic = source_run.get("italic", False)
        run.underline = source_run.get("underline", False)


def content_width_inches(profile: ExamProfile) -> float:
    total_twips = profile.width - profile.margins["left"] - profile.margins["right"]
    if profile.columns > 1:
        total_twips = (total_twips - profile.column_space * (profile.columns - 1)) / profile.columns
    return max(1.0, total_twips / 1440)


def image_display_width_inches(image: dict, profile: ExamProfile) -> float:
    max_width = content_width_inches(profile)
    try:
        from PIL import Image

        with Image.open(BytesIO(image["blob"])) as img:
            native_width = img.width / 96
            return min(max_width, max(1.0, native_width))
    except Exception:
        return max_width


def add_block_images(document: Document, profile: ExamProfile, images: list[dict]) -> None:
    for image in images:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run()
        try:
            run.add_picture(BytesIO(image["blob"]), width=Inches(image_display_width_inches(image, profile)))
        except Exception:
            fallback = paragraph.add_run(f"[圖片無法插入：{image.get('filename', 'unknown')}]")
            set_run_font(fallback, profile.font, max(8, profile.body_size - 1))


def insert_paragraph_xml(document: Document, paragraph_xml, profile: ExamProfile, images: list[dict] | None = None) -> Paragraph:
    rel_id_map = {}
    for image in images or []:
        old_rel_id = image.get("rel_id")
        if not old_rel_id:
            continue
        new_rel_id, _image = document.part.get_or_add_image(BytesIO(image["blob"]))
        rel_id_map[old_rel_id] = new_rel_id
    for blip in paragraph_xml.xpath(".//a:blip"):
        old_rel_id = blip.get(qn("r:embed"))
        if old_rel_id in rel_id_map:
            blip.set(qn("r:embed"), rel_id_map[old_rel_id])
    body = document._body._element
    sect_pr = body.sectPr
    if sect_pr is not None:
        body.insert(body.index(sect_pr), paragraph_xml)
    else:
        body.append(paragraph_xml)
    paragraph = Paragraph(paragraph_xml, document)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(profile.line_spacing)
    for run in paragraph.runs:
        if run.text:
            set_run_font(run, profile.font, profile.body_size)
    return paragraph


def insert_table_xml(document: Document, table_xml, profile: ExamProfile, images: dict | None = None) -> Table:
    for blip in table_xml.xpath(".//a:blip"):
        old_rel_id = blip.get(qn("r:embed"))
        if not old_rel_id or not images or old_rel_id not in images:
            continue
        new_rel_id, _image = document.part.get_or_add_image(BytesIO(images[old_rel_id]["blob"]))
        blip.set(qn("r:embed"), new_rel_id)
    body = document._body._element
    sect_pr = body.sectPr
    if sect_pr is not None:
        body.insert(body.index(sect_pr), table_xml)
    else:
        body.append(table_xml)
    table = Table(table_xml, document)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, profile.font, profile.table_font_size)
    return table


def set_paragraph_font(paragraph, font: str, size: int, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, font, size)
        run.bold = bold


def set_document_defaults(document: Document, profile: ExamProfile) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = profile.font
    normal.font.size = Pt(profile.body_size)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), profile.font)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(profile.line_spacing)


def set_section(profile: ExamProfile, document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Twips(profile.width)
    section.page_height = Twips(profile.height)
    section.top_margin = Twips(profile.margins["top"])
    section.right_margin = Twips(profile.margins["right"])
    section.bottom_margin = Twips(profile.margins["bottom"])
    section.left_margin = Twips(profile.margins["left"])
    section.header_distance = Twips(profile.header_footer[0])
    section.footer_distance = Twips(profile.header_footer[1])

    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    if not cols:
        sect_pr.append(cols_el)
    cols_el.set(qn("w:num"), str(profile.columns))
    cols_el.set(qn("w:space"), str(profile.column_space))
    if profile.column_separator:
        cols_el.set(qn("w:sep"), "1")
    elif qn("w:sep") in cols_el.attrib:
        del cols_el.attrib[qn("w:sep")]


def add_exam_header(document: Document, profile: ExamProfile, form: dict) -> None:
    school = form.get("school") or SCHOOL_NAME
    semester = form.get("semester") or "一百一十五學年度上學期"
    exam_name = form.get("exam_name") or "定期評量"
    subject = form.get("subject") or "科別"
    teacher = form.get("teacher") or "命題教師"
    grade = form.get("grade") or "年"
    class_name = form.get("class_name") or "班"
    duration = form.get("duration") or "時間"

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(school)
    set_run_font(r, profile.font, profile.title_size)
    r.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{semester}{exam_name}")
    set_run_font(r, profile.font, profile.title_size - 1)
    r.bold = True

    table = document.add_table(rows=2, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["科別", "命題教師", "班級", "座號", "姓名", "時間", "成績", "家長簽章"]
    values = [subject, teacher, f"{grade} {class_name}", "號", "", duration, "分", ""]
    for idx, text in enumerate(headers):
        set_cell_text(table.cell(0, idx), text, profile.font, 9, True)
    for idx, text in enumerate(values):
        set_cell_text(table.cell(1, idx), text, profile.font, 10)

    for row in table.rows:
        row.height = Twips(420)
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = OxmlElement("w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "60")
                node.set(qn("w:type"), "dxa")
                tc_mar.append(node)
            tc_pr.append(tc_mar)

    document.add_paragraph()


def add_body(document: Document, profile: ExamProfile, blocks: list[dict]) -> None:
    if not blocks:
        document.add_paragraph("請在此輸入試題內容。")
        return

    for block in blocks:
        if block["type"] == "paragraph":
            if block.get("preserve_xml") and block.get("xml") is not None:
                insert_paragraph_xml(document, deepcopy(block["xml"]), profile, block.get("images"))
                continue
            text = normalize_text(block.get("text", ""))
            if text:
                for line in text.splitlines():
                    text_line = normalize_text(line)
                    if not text_line:
                        continue
                    p = document.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(profile.line_spacing)
                    if looks_like_question(text_line):
                        p.paragraph_format.first_line_indent = Twips(-280)
                        p.paragraph_format.left_indent = Twips(280)
                        size = profile.question_size
                    else:
                        p.paragraph_format.first_line_indent = Twips(0)
                        size = profile.body_size
                    add_styled_runs(p, block.get("runs", []), text_line, profile.font, size)
            add_block_images(document, profile, block.get("images", []))
        elif block["type"] == "table":
            if block.get("xml") is not None:
                insert_table_xml(document, deepcopy(block["xml"]), profile, block.get("images"))
                continue
            rows = block["rows"]
            cols = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row_idx, row in enumerate(rows):
                for col_idx in range(cols):
                    text = row[col_idx] if col_idx < len(row) else ""
                    set_cell_text(table.cell(row_idx, col_idx), text, profile.font, profile.table_font_size, row_idx == 0)


def add_document_guards(document: Document) -> None:
    settings = document.settings.element
    for tag in ("w:doNotAutoFitConstrainedTables", "w:compat"):
        if not settings.xpath(f"./{tag}"):
            settings.append(OxmlElement(tag))


def export_docx_to_pdf(docx_path: Path) -> Path:
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return pdf_path
    except Exception:
        return convert_with_libreoffice(docx_path, ".pdf")


def count_pdf_pages(pdf_path: Path) -> int:
    from pypdf import PdfReader

    reader = PdfReader(str(pdf_path))
    return len(reader.pages)


def render_pdf_preview(pdf_path: Path) -> Path | None:
    try:
        import fitz  # type: ignore

        doc = fitz.open(str(pdf_path))
        if doc.page_count == 0:
            return None
        page = doc.load_page(0)
        pix = page.get_pixmap(matrix=fitz.Matrix(0.8, 0.8), alpha=False)
        preview_path = pdf_path.with_suffix(".page1.png")
        pix.save(str(preview_path))
        doc.close()
        return preview_path
    except Exception:
        return None


def compact_profile(profile: ExamProfile, level: int) -> ExamProfile:
    if level <= 0:
        return profile
    margin_cut = min(90 * level, 240)
    margins = {
        key: max(360, value - margin_cut)
        for key, value in profile.margins.items()
    }
    return replace(
        profile,
        margins=margins,
        body_size=max(9, profile.body_size - level),
        question_size=max(9, profile.question_size - level),
        table_font_size=max(8, profile.table_font_size - level),
        line_spacing=max(12, profile.line_spacing - level),
        column_space=max(320, profile.column_space - 25 * level),
    )


def build_report(profile: ExamProfile, stats: dict, pdf_pages: int | None, target_pages: int | None, compact_level: int) -> list[str]:
    report = [
        f"套用模板：{profile.label}",
        f"頁面尺寸：{profile.width} x {profile.height} twips",
        f"雙欄設定：{profile.columns} 欄，欄距 {profile.column_space} twips",
        f"內容統計：段落 {stats['paragraphs']}，表格 {stats['tables']}",
    ]
    if pdf_pages is not None:
        report.append(f"PDF 頁數：{pdf_pages} 頁")
    if target_pages:
        if pdf_pages is None:
            report.append("兩頁鎖定：未執行，因為 PDF 匯出失敗")
        elif pdf_pages <= target_pages:
            report.append(f"目標頁數：已符合 {target_pages} 頁以內")
        else:
            report.append(f"目標頁數：仍超過 {target_pages} 頁，建議檢查表格、圖片或長題組")
    if compact_level:
        report.append(f"自動壓縮：已套用第 {compact_level} 級壓縮")
    return report


def build_exam_docx(source: Path, output: Path, profile: ExamProfile, form: dict) -> dict:
    source_for_read = convert_doc_to_docx(source) if source.suffix.lower() == ".doc" else source
    blocks = iter_docx_blocks(source_for_read)
    document = Document()
    set_document_defaults(document, profile)
    set_section(profile, document)
    add_exam_header(document, profile, form)
    add_body(document, profile, blocks)
    add_document_guards(document)
    document.save(str(output))
    return {
        "paragraphs": sum(1 for b in blocks if b["type"] == "paragraph"),
        "tables": sum(1 for b in blocks if b["type"] == "table"),
        "formulas": sum(1 for b in blocks if b.get("has_formula")),
        "ruby": sum(1 for b in blocks if b.get("has_ruby")),
    }


def build_with_page_check(source: Path, output: Path, profile: ExamProfile, form: dict, target_pages: int | None = None) -> dict:
    last_error = None
    max_level = 3 if target_pages else 0
    for level in range(max_level + 1):
        candidate_profile = compact_profile(profile, level)
        stats = build_exam_docx(source, output, candidate_profile, form)
        pdf_path = None
        pdf_pages = None
        preview_path = None
        try:
            pdf_path = export_docx_to_pdf(output)
            pdf_pages = count_pdf_pages(pdf_path)
            preview_path = render_pdf_preview(pdf_path)
        except Exception as exc:
            last_error = str(exc)
        if not target_pages or pdf_pages is None or pdf_pages <= target_pages or level == max_level:
            return {
                "stats": stats,
                "pdf": pdf_path.name if pdf_path else None,
                "pdf_pages": pdf_pages,
                "preview": preview_path.name if preview_path else None,
                "compact_level": level,
                "warnings": [last_error] if last_error else [],
                "report": build_report(candidate_profile, stats, pdf_pages, target_pages, level),
            }
    raise RuntimeError("無法完成轉換。")


@app.get("/")
def index():
    site_url = os.environ.get("PUBLIC_SITE_URL") or request.url_root
    if not site_url.endswith("/"):
        site_url += "/"
    return render_template(
        "index.html",
        profiles=PROFILES,
        app_title=APP_TITLE,
        app_description=APP_DESCRIPTION,
        school_name=SCHOOL_NAME,
        site_url=site_url,
        asset_version=ASSET_VERSION,
        app_version=load_app_version(),
    )


@app.get("/sw.js")
def service_worker():
    response = send_file(ROOT / "sw.js", mimetype="application/javascript")
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    return response


@app.get("/version.json")
def version_json():
    response = send_file(VERSION_PATH, mimetype="application/json")
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    return response


@app.post("/api/analyze")
def analyze():
    if request.method == "OPTIONS":
        return ("", 204)
    ensure_work_dirs()
    file = request.files.get("file")
    if not file or not allowed_file(file.filename):
        return error_response("請上傳 .doc 或 .docx 檔案。", 400, "INVALID_FILE")
    job_id, upload_dir, _generated_dir = create_job_dirs()
    upload_path = upload_dir / safe_upload_name(file.filename)
    file.save(upload_path)
    try:
        read_path = convert_doc_to_docx(upload_path) if upload_path.suffix.lower() == ".doc" else upload_path
        blocks = iter_docx_blocks(read_path)
        preview = preview_blocks(blocks)
        return jsonify({"ok": True, "job_id": job_id, "expires_in_seconds": JOB_TTL_SECONDS, "filename": file.filename, "blocks": len(blocks), "preview": preview})
    except Exception as exc:
        return error_response("分析失敗，請確認 Word 檔沒有加密或損毀。", 400, "ANALYZE_FAILED", [str(exc)])


@app.post("/api/convert")
def convert():
    if request.method == "OPTIONS":
        return ("", 204)
    ensure_work_dirs()
    file = request.files.get("file")
    profile_key = request.form.get("profile", "b4-horizontal")
    if profile_key not in PROFILES:
        return error_response("找不到指定的格式設定，請重新選擇 B4/A4 或直式/橫式模板。", 400, "INVALID_PROFILE")
    if not file or not allowed_file(file.filename):
        return error_response("請上傳 .doc 或 .docx 檔案。", 400, "INVALID_FILE")

    job_id, upload_dir, generated_dir = create_job_dirs()
    upload_path = upload_dir / safe_upload_name(file.filename)
    file.save(upload_path)
    out_name = f"{Path(file.filename).stem}-格式校正-{uuid.uuid4().hex[:6]}.docx"
    output_path = generated_dir / out_name
    target_pages_raw = request.form.get("target_pages", "").strip()
    target_pages = int(target_pages_raw) if target_pages_raw.isdigit() else None

    try:
        result = build_with_page_check(upload_path, output_path, PROFILES[profile_key], request.form, target_pages)
    except Exception as exc:
        return error_response("轉換失敗。請確認檔案沒有加密、損毀，或嘗試先移除大型圖片後再上傳。", 400, "CONVERT_FAILED", [str(exc)])

    return jsonify(
        {
            "ok": True,
            "job_id": job_id,
            "expires_in_seconds": JOB_TTL_SECONDS,
            "download": public_file_path("download", job_id, out_name),
            "pdf": public_file_path("download", job_id, result["pdf"]) if result["pdf"] else None,
            "preview": public_file_path("preview", job_id, result["preview"]) if result["preview"] else None,
            "stats": result["stats"],
            "pdf_pages": result["pdf_pages"],
            "compact_level": result["compact_level"],
            "warnings": result["warnings"],
            "report": result["report"],
        }
    )


@app.get("/download/<job_id>/<path:filename>")
def download(job_id: str, filename: str):
    path = resolve_job_file(GENERATED_DIR, job_id, filename)
    if not path:
        return error_response("檔案不存在或下載連結已過期，請重新轉換一次。", 404, "FILE_EXPIRED")
    return send_file(path, as_attachment=True, download_name=filename)


@app.get("/preview/<job_id>/<path:filename>")
def preview_file(job_id: str, filename: str):
    path = resolve_job_file(GENERATED_DIR, job_id, filename)
    if not path or path.suffix.lower() != ".png":
        return error_response("預覽圖不存在或已過期，請重新轉換一次。", 404, "PREVIEW_EXPIRED")
    return send_file(path, mimetype="image/png")


@app.post("/api/clear")
def clear_generated():
    for directory in (UPLOAD_DIR, GENERATED_DIR):
        shutil.rmtree(directory, ignore_errors=True)
        directory.mkdir(parents=True, exist_ok=True)
    return jsonify({"ok": True})


if __name__ == "__main__":
    ensure_work_dirs()
    port = int(os.environ.get("PORT", "5127"))
    debug = os.environ.get("FLASK_DEBUG", "1") == "1"
    app.run(host="0.0.0.0", port=port, debug=debug)
